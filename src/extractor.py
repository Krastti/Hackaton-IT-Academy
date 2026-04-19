import json
import importlib
import logging
import os
import re
import shutil
import subprocess
from dataclasses import dataclass, field
from threading import Lock
from typing import Any, Callable, ClassVar, Dict, List, Literal, Optional, Sequence
from contextlib import redirect_stdout, redirect_stderr

import cv2
import docx
import pandas as pd
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text
from PIL import Image

logger = logging.getLogger(__name__)

ExtractStatus = Literal['success', 'empty', 'unsupported', 'failed', 'partial']
TableFrame = Any

"""Неизменяемый контейнер параметров, управляющих поведением экстрактора"""


@dataclass(frozen=True)
class ExtractConfig:
    enable_ocr: bool = True
    image_rotate_passes: int = 3  # Оставлено для совместимости, Surya отлично читает под углом
    video_max_seconds: float = 120.0
    video_fps_sample: float = 1.0
    max_pdf_pages: int = 200
    max_table_rows: int = 1000
    max_json_chars: int = 200000
    max_chars: int = 200000


@dataclass
class ExtractResult:
    file_path: str
    text: str
    status: ExtractStatus
    extractor: str
    warnings: List[str] = field(default_factory=list)
    error: Optional[str] = None


@dataclass
class ExtractContext:
    config: ExtractConfig
    warnings: List[str] = field(default_factory=list)
    handler_name: str = ''


class FileExtractor:
    _default_config: ClassVar[ExtractConfig] = ExtractConfig()
    _pdf_module: ClassVar[Optional[Any]] = None
    _pdf_module_lock: ClassVar[Lock] = Lock()
    _pdf_module_checked: ClassVar[bool] = False

    # Ленивая загрузка Surya OCR
    _surya_foundation = None
    _surya_recognition = None
    _surya_detection = None
    _surya_lock: ClassVar[Lock] = Lock()

    _handlers: ClassVar[Dict[str, str]] = {
        '.pdf': '_extract_pdf',
        '.docx': '_extract_docx',
        '.doc': '_extract_doc',
        '.csv': '_extract_table', '.tsv': '_extract_table',
        '.parquet': '_extract_table', '.xls': '_extract_table', '.xlsx': '_extract_table',
        '.json': '_extract_json', '.rtf': '_extract_rtf',
        '.html': '_extract_html', '.htm': '_extract_html', '.xml': '_extract_html',
        '.png': '_extract_image', '.jpg': '_extract_image', '.jpeg': '_extract_image',
        '.tif': '_extract_image', '.tiff': '_extract_image', '.gif': '_extract_image',
        '.bmp': '_extract_image', '.webp': '_extract_image',
        '.mp4': '_extract_video', '.avi': '_extract_video', '.mov': '_extract_video', '.mkv': '_extract_video',
        '.txt': '_extract_plain_text', '.md': '_extract_plain_text', '.log': '_extract_plain_text',
    }

    @classmethod
    def _get_surya_predictors(cls):
        """Потокобезопасная ленивая инициализация тяжелых моделей Surya"""
        if cls._surya_recognition is None:
            with cls._surya_lock:
                if cls._surya_recognition is None:
                    logger.info("Загрузка моделей Surya OCR...")
                    from surya.foundation import FoundationPredictor
                    from surya.recognition import RecognitionPredictor
                    from surya.detection import DetectionPredictor

                    cls._surya_foundation = FoundationPredictor()
                    cls._surya_recognition = RecognitionPredictor(cls._surya_foundation)
                    cls._surya_detection = DetectionPredictor()
                    logger.info("Модели Surya OCR успешно загружены.")
        return cls._surya_recognition, cls._surya_detection

    @classmethod
    def _get_pdf_module(cls) -> Optional[Any]:
        if cls._pdf_module_checked:
            return cls._pdf_module

        with cls._pdf_module_lock:
            if cls._pdf_module_checked:
                return cls._pdf_module

            for module_name in ('pymupdf', 'fitz'):
                try:
                    module = importlib.import_module(module_name)
                except (ModuleNotFoundError, ImportError):
                    continue
                if hasattr(module, 'open'):
                    cls._pdf_module = module
                    cls._pdf_module_checked = True
                    return module

            cls._pdf_module_checked = True
        return None

    @classmethod
    def extract(cls, file_path: str, config: Optional[ExtractConfig] = None) -> ExtractResult:
        ext: str = os.path.splitext(file_path)[1].lower()
        handler_name = cls._handlers.get(ext)
        active_config = config or cls._default_config
        context = ExtractContext(config=active_config)

        if handler_name is None:
            warning = f'Неподдерживаемое расширение файла: {ext or "<none>"}'
            logger.warning('%s (%s)', warning, file_path)
            return ExtractResult(file_path=file_path, text='', status='unsupported', extractor='unsupported',
                                 warnings=[warning])

        extractor = cls()
        context.handler_name = handler_name
        handler = getattr(extractor, handler_name)

        try:
            text = handler(file_path, context)
        except Exception as exc:
            logger.exception('Экстрактор %s завершился с ошибкой для %s', handler_name, os.path.basename(file_path))
            return ExtractResult(file_path=file_path, text='', status='failed', extractor=handler_name,
                                 warnings=context.warnings, error=str(exc))

        normalized_text = cls._finalize_text(text, context)
        status = cls._resolve_status(normalized_text, context.warnings)

        if status in ('failed', 'partial'):
            logger.warning('Экстрактор %s завершился со статусом %s для %s: %s', handler_name, status,
                           os.path.basename(file_path), '; '.join(context.warnings))
        elif status == 'empty':
            logger.info('Экстрактор %s вернул пустой текст для %s', handler_name, os.path.basename(file_path))

        return ExtractResult(file_path=file_path, text=normalized_text, status=status, extractor=handler_name,
                             warnings=context.warnings)

    @classmethod
    def extract_text(cls, file_path: str, config: Optional[ExtractConfig] = None) -> str:
        return cls.extract(file_path, config=config).text

    @staticmethod
    def _resolve_status(text: str, warnings: Sequence[str]) -> ExtractStatus:
        if not text.strip():
            return 'partial' if warnings else 'empty'
        return 'partial' if warnings else 'success'

    @staticmethod
    def _finalize_text(text: str, context: ExtractContext) -> str:
        collapsed = ' '.join(text.split())
        if len(collapsed) > context.config.max_chars:
            context.warnings.append(f'Текст обрезан до {context.config.max_chars} символов')
            return collapsed[: context.config.max_chars]
        return collapsed

    @staticmethod
    def _normalize_block(value: Any) -> str:
        return ' '.join(str(value).split())

    @classmethod
    def _merge_blocks(cls, blocks: Sequence[str]) -> str:
        unique_blocks: List[str] = []
        seen: set[str] = set()
        for block in blocks:
            normalized = cls._normalize_block(block)
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            unique_blocks.append(normalized)
        return ' '.join(unique_blocks)

    @staticmethod
    def _truncate_text(text: str, limit: int, context: ExtractContext, label: str) -> str:
        if len(text) > limit:
            context.warnings.append(f'{label} обрезан до {limit} символов')
            return text[:limit]
        return text

    # --- ДОКУМЕНТЫ ---
    def _extract_pdf(self, file_path: str, context: ExtractContext) -> str:
        pdf_module = self._get_pdf_module()
        if pdf_module is None:
            context.warnings.append('PDF-обработка недоступна: не найден модуль PyMuPDF/fitz')
            return ''

        text_blocks: List[str] = []
        with pdf_module.open(file_path) as doc:
            page_limit = min(len(doc), context.config.max_pdf_pages)
            if len(doc) > page_limit:
                context.warnings.append(f'PDF ограничен первыми {page_limit} страницами из {len(doc)}')

            for page_index in range(page_limit):
                page = doc[page_index]
                page_text = self._normalize_block(page.get_text())
                if page_text:
                    text_blocks.append(page_text)
                    continue

                ocr_text = self._extract_pdf_page_with_ocr(page, context)
                if ocr_text:
                    text_blocks.append(ocr_text)

        return self._merge_blocks(text_blocks)

    def _extract_docx(self, file_path: str, context: ExtractContext) -> str:
        doc = docx.Document(file_path)
        parts: List[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend([cell.text for cell in row.cells])
        return self._merge_blocks(parts)

    def _extract_doc(self, file_path: str, context: ExtractContext) -> str:
        antiword_path = shutil.which('antiword')
        if antiword_path is None:
            fallback_text = self._extract_doc_as_binary_text(file_path, context)
            if fallback_text:
                context.warnings.append('antiword не найден; использовано извлечение из бинарного .doc')
                return fallback_text
            context.warnings.append('Для .doc требуется antiword')
            return ''

        completed = subprocess.run([antiword_path, file_path], capture_output=True, text=True, encoding='utf-8',
                                   errors='ignore', check=False)
        if completed.returncode != 0:
            raise RuntimeError(completed.stderr.strip() or 'antiword завершился с ошибкой')
        return completed.stdout

    def _extract_rtf(self, file_path: str, context: ExtractContext) -> str:
        raw_text = self._read_text_with_fallback(file_path, context)
        return str(rtf_to_text(raw_text)) if raw_text else ''

    def _extract_table(self, file_path: str, context: ExtractContext) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext in ('.csv', '.tsv'):
            df = self._read_csv_with_fallback(file_path, context, delimiter='\t' if ext == '.tsv' else ',')
        elif ext == '.parquet':
            df = pd.read_parquet(file_path)
        else:
            df = pd.read_excel(file_path)
        return self._serialize_table(df, context)

    def _read_csv_with_fallback(self, file_path: str, context: ExtractContext, delimiter: str = ',') -> TableFrame:
        last_error = None
        for enc in ('utf-8-sig', 'utf-8', 'cp1251', 'latin-1'):
            try:
                return pd.read_csv(file_path, on_bad_lines='skip', encoding=enc, sep=delimiter, low_memory=False)
            except UnicodeDecodeError as exc:
                last_error = str(exc)
                continue
        raise UnicodeDecodeError('csv', b'', 0, 1, last_error or 'Unknown Error')

    def _extract_doc_as_binary_text(self, file_path: str, context: ExtractContext) -> str:
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
        except OSError as exc:
            context.warnings.append(f'Ошибка чтения .doc: {exc}')
            return ''

        candidates = []
        for enc in ('utf-16le', 'cp1251', 'latin-1'):
            try:
                decoded = raw_data.decode(enc, errors='ignore')
                candidates.extend(re.findall(r'[\wА-Яа-яЁё.,;:()@/\- ]{4,}', decoded))
            except LookupError:
                continue
        return self._merge_blocks([c.strip() for c in candidates if c.strip()])

    def _serialize_table(self, df: TableFrame, context: ExtractContext) -> str:
        if df.empty: return ''
        limited_df = df.head(context.config.max_table_rows).fillna('')
        if len(df.index) > len(limited_df.index):
            context.warnings.append(f'Таблица ограничена {len(limited_df.index)} строками')

        cols = [str(c) for c in limited_df.columns]
        lines = [f'Столбцы: {", ".join(cols)}']
        for _, row in limited_df.iterrows():
            parts = [f'{c}: {self._normalize_block(row[c])}' for c in cols if self._normalize_block(row[c])]
            if parts: lines.append(' | '.join(parts))
        return '\n'.join(lines)

    def _extract_json(self, file_path: str, context: ExtractContext) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            data = json.load(f)
        return self._truncate_text(json.dumps(data, ensure_ascii=False, indent=2), context.config.max_json_chars,
                                   context, 'JSON')

    def _extract_html(self, file_path: str, context: ExtractContext) -> str:
        raw_text = self._read_text_with_fallback(file_path, context)
        return BeautifulSoup(raw_text, 'html.parser').get_text(separator=' ', strip=True) if raw_text else ''

    def _extract_plain_text(self, file_path: str, context: ExtractContext) -> str:
        return self._read_text_with_fallback(file_path, context)

    def _read_text_with_fallback(self, file_path: str, context: ExtractContext) -> str:
        for enc in ('utf-8-sig', 'utf-8', 'cp1251', 'latin-1'):
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        context.warnings.append('Не удалось декодировать текстовый файл')
        return ''

    # --- МЕДИА (OCR: Surya + Батчинг) ---
    def _run_surya_batch(self, images: List[Image.Image], context: ExtractContext) -> List[str]:
        """Обрабатывает список PIL изображений через Surya (Тихо и безопасно)"""
        if not context.config.enable_ocr or not images:
            return []

        rec_pred, det_pred = self._get_surya_predictors()
        text_blocks = []

        try:
            with open(os.devnull, 'w') as fnull:
                with redirect_stdout(fnull), redirect_stderr(fnull):
                    predictions = rec_pred(images, det_predictor=det_pred)

            for p in predictions:
                if p and p.text_lines:
                    text_blocks.extend([line.text for line in p.text_lines])
        except Exception as exc:
            context.warnings.append(f"Ошибка Surya OCR: {exc}")

        return text_blocks

    def _extract_pdf_page_with_ocr(self, page: Any, context: ExtractContext) -> str:
        if not hasattr(page, 'get_pixmap'): return ''
        try:
            pixmap = page.get_pixmap(dpi=200)
            mode = "RGBA" if pixmap.alpha else "RGB"
            samples = pixmap.samples if hasattr(pixmap, 'samples') else pixmap.samples_mv

            pil_image = Image.frombytes(mode, [pixmap.width, pixmap.height], samples)
            if mode == "RGBA":
                bg = Image.new("RGB", pil_image.size, (255, 255, 255))
                bg.paste(pil_image, mask=pil_image.split()[3])
                pil_image = bg

            pil_image.thumbnail((2048, 2048))  # Защита от OOM
        except Exception as exc:
            context.warnings.append(f'Ошибка рендеринга страницы PDF: {exc}')
            return ''

        blocks = self._run_surya_batch([pil_image], context)
        return self._merge_blocks(blocks)

    def _extract_image(self, file_path: str, context: ExtractContext) -> str:
        try:
            pil_image = Image.open(file_path).convert("RGB")
            pil_image.thumbnail((2048, 2048))  # Ускорение работы трансформера
        except Exception as exc:
            context.warnings.append(f"Ошибка чтения изображения {exc}")
            return ''

        blocks = self._run_surya_batch([pil_image], context)
        return self._merge_blocks(blocks)

    def _extract_video(self, file_path: str, context: ExtractContext) -> str:
        if not context.config.enable_ocr:
            context.warnings.append('OCR для видео отключён в конфигурации')
            return ''

        capture = cv2.VideoCapture(file_path)
        if not capture.isOpened():
            raise ValueError('OpenCV не удалось открыть видео')

        fps_raw = capture.get(cv2.CAP_PROP_FPS)
        fps_value = float(fps_raw) if fps_raw and fps_raw > 0 else 24.0
        total_frames = int(capture.get(cv2.CAP_PROP_FRAME_COUNT))
        total_seconds = total_frames / fps_value if fps_value > 0 else 0.0
        max_seconds = min(total_seconds, context.config.video_max_seconds)
        frame_step = max(int(round(fps_value / max(context.config.video_fps_sample, 0.1))), 1)

        batch_images = []
        all_text_blocks = []
        frame_index = 0

        try:
            if total_seconds > context.config.video_max_seconds:
                context.warnings.append(
                    f'OCR видео ограничен первыми {context.config.video_max_seconds:.1f} сек. из {total_seconds:.1f}'
                )

            while True:
                capture.set(cv2.CAP_PROP_POS_FRAMES, frame_index)
                success, frame = capture.read()
                if not success: break

                if (frame_index / fps_value if fps_value > 0 else 0.0) > max_seconds: break

                # Конвертация кадра в PIL и ресайз для дикого ускорения
                frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                pil_img = Image.fromarray(frame_rgb)
                pil_img.thumbnail((1024, 1024))

                batch_images.append(pil_img)

                # Отправляем пачку кадров в Surya разом
                if len(batch_images) >= 4:
                    all_text_blocks.extend(self._run_surya_batch(batch_images, context))
                    batch_images = []

                frame_index += frame_step

            # Добиваем остатки кадров
            if batch_images:
                all_text_blocks.extend(self._run_surya_batch(batch_images, context))

        finally:
            capture.release()

        return self._merge_blocks(all_text_blocks)